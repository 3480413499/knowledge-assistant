"""Tests for file_parser.py"""
import pytest
from pathlib import Path
from src.file_parser import FileParser, FileParseError, UnsupportedFormatError
from unittest.mock import patch


class TestFileParser:
    def test_parse_txt_returns_text(self, tmp_path):
        """解析 TXT 文件应返回文本内容"""
        f = tmp_path / "test.txt"
        f.write_text("你好，这是一个测试文档。\n第二行内容。", encoding="utf-8")
        result = FileParser.parse(f)
        assert "测试文档" in result
        assert "第二行" in result

    def test_parse_unsupported_extension_raises(self, tmp_path):
        """不支持的扩展名应抛出 UnsupportedFormatError"""
        f = tmp_path / "test.xyz"
        f.write_text("content")
        with pytest.raises(UnsupportedFormatError) as exc:
            FileParser.parse(f)
        assert "xyz" in str(exc.value)

    def test_file_too_large_raises(self, tmp_path):
        """超过 50MB 文件应抛出 FileParseError"""
        f = tmp_path / "large.txt"
        f.write_text("small")
        with patch("pathlib.Path.stat") as mock_stat:
            mock_stat.return_value.st_size = 51 * 1024 * 1024
            with pytest.raises(FileParseError) as exc:
                FileParser.parse(f)
            assert "50MB" in str(exc.value)

    def test_txt_empty_file_returns_empty_string(self, tmp_path):
        """空文件应返回空字符串"""
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        result = FileParser.parse(f)
        assert result == ""

    def test_parse_pdf_with_text(self, tmp_path):
        """解析含文字的 PDF 应提取出文本"""
        import fitz
        doc = fitz.open()
        # Use a CJK-supporting font name so Chinese renders correctly
        doc.new_page().insert_text((72, 72), "Hello PDF 测试内容", fontname="china-s")
        pdf_path = tmp_path / "test_text.pdf"
        doc.save(str(pdf_path))
        doc.close()
        result = FileParser.parse(pdf_path)
        assert "测试内容" in result

    def test_pdf_no_text_detected_raises(self, tmp_path):
        """扫描版 PDF（无文字层）应抛出 FileParseError"""
        import fitz
        doc = fitz.open()
        doc.new_page()  # add page with no text
        doc.save(str(tmp_path / "scan.pdf"))
        doc.close()
        with pytest.raises(FileParseError) as exc:
            FileParser.parse(tmp_path / "scan.pdf")
        assert "OCR" in str(exc.value) or "文字" in str(exc.value)

    def test_parse_docx(self, tmp_path):
        """解析 DOCX 文件"""
        from docx import Document
        doc = Document()
        doc.add_paragraph("这是Word文档内容。")
        doc.add_paragraph("第二个段落。")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        result = FileParser.parse(path)
        assert "Word文档内容" in result
        assert "第二个段落" in result

    def test_compute_md5(self, tmp_path):
        """计算文件 MD5 指纹"""
        f = tmp_path / "md5test.txt"
        f.write_text("hello world", encoding="utf-8")
        md5 = FileParser.compute_md5(f)
        import hashlib
        expected = hashlib.md5(b"hello world").hexdigest()
        assert md5 == expected
        assert len(md5) == 32
